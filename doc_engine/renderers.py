import os
import logging
from docx import Document

from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .config import IMAGE_DOCX_WIDTH_INCHES

logger = logging.getLogger(__name__)


class MarkdownRenderer:
    def render(self, flat_model, out_path, context):
        lines = []
        lines.append("<!-- AUTO-GENERATED START -->\n")
        lines.append(f"# {context['title']}\n")
        lines.append(f"**Application Name:** {context['app_name']}\n")
        lines.append(f"**Guide Type:** {context['guide_type']}\n")
        lines.append(f"**{context['lbl_generated']}:** {context['date']}  \n")
        lines.append(f"**{context['lbl_language']}:** {context['language']}  \n")
        lines.append(f"**{context['lbl_theme']}:** {context['theme']}  \n")
        lines.append(f"**{context['lbl_device']}:** {context['device']}  \n")
        lines.append(f"**{context['lbl_version']}:** {context['version']}  \n")
        lines.append(f"**{context['lbl_generated_by']}:** WealthFlow Documentation Engine  \n")
        lines.append("\n---\n")
        
        lines.append(f"## {context['toc_title']}\n")
        for item in flat_model:
            anchor = item.title.lower().replace(' ', '-')
            indent = "  " * (len(item.hierarchical_number.split('.')) - 1)
            lines.append(f"{indent}- **{item.hierarchical_number}** [{item.title}](#{anchor})\n")
        lines.append("\n---\n")

        for i, item in enumerate(flat_model, 1):
            h_level = min(2 + len(item.hierarchical_number.split('.')) - 1, 6)
            h_prefix = "#" * h_level
            lines.append(f"{h_prefix} {item.hierarchical_number} {item.title}\n")
            
            nav_path = " &rarr; ".join(item.navigation)
            lines.append(f"**{context['nav_title']}:** {nav_path}\n")
            
            if item.purpose:
                lines.append(f"**{context['purpose_title']}:** {item.purpose}\n")
            
            for s_idx, s_path in enumerate(item.screenshots):
                if os.path.exists(s_path):
                    lines.append("<figure>")
                    lines.append(f"<img src=\"file:///{s_path.replace(chr(92), '/')}\" style=\"max-width: 100%; height: auto; display: block; margin: 0 auto;\" alt=\"{item.title}\">")
                    lines.append(f"<figcaption style=\"text-align:center; font-style:italic;\">{context['figure_title']} {i}.{s_idx+1}: {item.title}</figcaption>")
                    lines.append("</figure>\n")
            
            if context['is_technical']:
                lines.append(f"**{context['tech_notes_title']}:**\n")
                lines.append(f"- Route Hierarchy: `{nav_path}`\n")
                lines.append(f"- Base Route: `{getattr(item, 'route', 'N/A')}`\n")
                lines.append("\n")
            elif item.steps:
                lines.append(f"**{context['steps_title']}:**\n")
                for s_idx, step in enumerate(item.steps, 1):
                    lines.append(f"{s_idx}. {step}")
                lines.append("\n")
                
            if item.siblings:
                lines.append("**Related Pages:**")
                for s_idx, sib in enumerate(item.siblings, 1):
                    lines.append(f"- {sib}")
                lines.append("\n")
            lines.append("---\n")
            
        lines.append("<!-- AUTO-GENERATED END -->\n")
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

class HtmlRenderer:

    def render(self, md_path, html_path, context):
        with open(md_path, "r", encoding="utf-8") as f:
            md_text = f.read()
            
        md_text = md_text.replace("<!-- AUTO-GENERATED START -->\n", "").replace("<!-- AUTO-GENERATED END -->\n", "")
        import markdown
        html_content = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'toc'])
        
        is_dark = context.get('theme', '').lower() == 'dark'
        
        if is_dark:
            bg_color = "#121212"
            container_bg = "#1e1e1e"
            text_color = "#e0e0e0"
            heading_color = "#ffffff"
            hr_color = "#333333"
            code_bg = "#2d2d2d"
            caption_color = "#aaaaaa"
        else:
            bg_color = "#f4f7f6"
            container_bg = "#ffffff"
            text_color = "#333333"
            heading_color = "#2c3e50"
            hr_color = "#eeeeee"
            code_bg = "#f4f4f4"
            caption_color = "#666666"

        final_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset='utf-8'>
<title>{context['title']}</title>
<style>
body, html {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; line-height: 1.6; color: {text_color}; background-color: {bg_color}; max-width: 1200px; margin: 0 auto; padding: 2rem; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
h1, h2, h3, h4, h5, h6 {{ color: {heading_color}; }}
img {{ max-width: 100%; height: auto; display: block; margin: 2rem auto; border-radius: 8px; box-shadow: 0 4px 6px rgba(0,0,0,0.3); }}
figcaption {{ text-align: center; font-style: italic; color: {caption_color}; margin-top: 0.5rem; }}
.container {{ background: {container_bg}; padding: 2rem; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
hr {{ border: 0; height: 1px; background: {hr_color}; margin: 2rem 0; }}
code {{ background: {code_bg}; padding: 0.2rem 0.4rem; border-radius: 4px; font-family: monospace; }}
table {{ width: 100%; border-collapse: collapse; margin: 1rem 0; }}
th, td {{ padding: 0.75rem; border: 1px solid {hr_color}; text-align: left; }}
th {{ background-color: {code_bg}; }}
a {{ color: #3498db; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
</style>
</head>
<body><div class='container'>{html_content}</div></body></html>"""
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(final_html)

class PdfRenderer:
    """
    Renders HTML documentation into PDF format.
    Delegates PDF rendering to the active PlaywrightBackend strategy.
    """
    def render(self, html_path, pdf_path):
        from .playwright_engine import get_playwright_backend
        backend = get_playwright_backend()
        success = backend.render_pdf(html_path, pdf_path)
        if not success:
            logger.error(f"PDF generation failed using {backend.__class__.__name__}")



class DocxRenderer:
    def _add_page_number(self, run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def render(self, flat_model, out_path, context):
        doc = Document()
        is_dark = context.get('theme', '').lower() == 'dark'
        if is_dark:
            # Set background color
            background = OxmlElement('w:background')
            background.set(qn('w:color'), '121212')
            doc.element.insert(0, background)
            
            if not doc.settings.element.xpath('./w:displayBackgroundShape'):
                disp = OxmlElement('w:displayBackgroundShape')
                doc.settings.element.append(disp)
                
            # Set text colors for styles
            for style_name in ['Normal', 'List Paragraph', 'List Bullet']:
                try:
                    style = doc.styles[style_name]
                    if hasattr(style, 'font'):
                        style.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
                except KeyError:
                    pass
            for i in range(1, 4):
                try:
                    style = doc.styles[f'Heading {i}']
                    if hasattr(style, 'font'):
                        style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                except KeyError:
                    pass
            try:
                doc.styles['Title'].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            except KeyError:
                pass
        section = doc.sections[0]
        header_p = section.header.paragraphs[0]
        header_p.text = context['title']
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(context['page_title'] + " ") 
        self._add_page_number(run)

        doc.add_heading(f"{context['app_name']}", 0)
        doc.add_heading(f"{context['guide_type']}", 1)
        p = doc.add_paragraph()
        p.add_run(f"{context['lbl_generated']}: {context['date']}\n")
        p.add_run(f"{context['lbl_language']}: {context['language']}\n")
        p.add_run(f"{context['lbl_theme']}: {context['theme']}\n")
        p.add_run(f"{context['lbl_device']}: {context['device']}\n")
        p.add_run(f"{context['lbl_version']}: {context['version']}\n")
        p.add_run(f"{context['lbl_generated_by']}: WealthFlow Documentation Engine")
        doc.add_page_break()
        
        doc.add_heading(context['toc_title'], level=1)
        for i, item in enumerate(flat_model, 1):
            indent_level = len(item.hierarchical_number.split('.')) - 1
            p_toc = doc.add_paragraph(f"{item.hierarchical_number} {item.title}")
            p_toc.paragraph_format.left_indent = Inches(0.3 * indent_level)
        doc.add_page_break()
        
        for i, item in enumerate(flat_model, 1):
            h_level = min(1 + len(item.hierarchical_number.split('.')) - 1, 3)
            doc.add_heading(f"{item.hierarchical_number} {item.title}", level=h_level)
            
            nav_path = " > ".join(item.navigation)
            p = doc.add_paragraph()
            p.add_run(f"{context['nav_title']}: ").bold = True
            p.add_run(nav_path)
            
            if item.purpose:
                p2 = doc.add_paragraph()
                p2.add_run(f"{context['purpose_title']}: ").bold = True
                p2.add_run(item.purpose)
            
            if context['is_technical']:
                doc.add_paragraph(context['tech_notes_title'] + ":", style='Heading 3')
                doc.add_paragraph(f"Route Hierarchy: {nav_path}", style='List Bullet')
                doc.add_paragraph(f"Base Route: {getattr(item, 'route', 'N/A')}", style='List Bullet')
            elif item.steps:
                doc.add_paragraph(context['steps_title'] + ":", style='Heading 3')
                for s_idx, step in enumerate(item.steps, 1):
                    doc.add_paragraph(f"{s_idx}. {step}", style='List Paragraph')
                
            for s_idx, s_path in enumerate(item.screenshots):
                if os.path.exists(s_path):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(s_path, width=Inches(IMAGE_DOCX_WIDTH_INCHES))
                        
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_cap = p_cap.add_run(f"{context['figure_title']} {i}.{s_idx+1}: {item.title}")
                        r_cap.italic = True
                    except Exception as e:
                        logger.warning(f"Could not add image {s_path} to docx: {e}")
                        
            if item.siblings:
                doc.add_paragraph("Related Pages:", style='Heading 3')
                for sib in item.siblings:
                    doc.add_paragraph(sib, style='List Bullet')
        doc.save(out_path)
