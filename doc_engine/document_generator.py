import os
import json
import logging
import re
import shutil
from datetime import datetime
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import django
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "wealthflow.settings")
django.setup()

from django.conf import settings
from django.utils.translation import gettext as _
from django.utils import translation
import subprocess

logger = logging.getLogger(__name__)

DOC_ENGINE_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = getattr(settings, 'BASE_DIR', os.path.dirname(DOC_ENGINE_DIR))
DOCS_DIR = os.path.join(BASE_DIR, "docs")
GENERATED_DIR = os.path.join(DOCS_DIR, "generated")
MANIFEST_FILE = os.path.join(GENERATED_DIR, "manifest.json")
CONTENT_FILE = os.path.join(DOC_ENGINE_DIR, "content", "page_descriptions.json")
TEMPLATES_DIR = os.path.join(DOC_ENGINE_DIR, "templates")

class DocumentationGenerator:
    def __init__(self):
        self.manifest = self._load_json(MANIFEST_FILE, default={"pages": []})
        self.content = self._load_json(CONTENT_FILE, default={})
        self.internal_model = []
        
        # Validation report
        self.validation_errors = []
        self.validation_warnings = []
        self.timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.output_base_dir = os.path.join(GENERATED_DIR, self.timestamp)
        self.latest_symlink = os.path.join(GENERATED_DIR, "latest")
        
        lang_code = self.manifest.get('language', 'en').lower()
        if lang_code == 'en':
            lang_code = 'en-us' # Django default
        translation.activate(lang_code)
        
        self._validate_manifest()
        self._build_internal_model()

    def _load_json(self, path, default=None):
        if not os.path.exists(path):
            return default
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Failed to load {path}: {e}")
            return default

    def _get_stable_key(self, page_entry):
        keys = [page_entry.get('route')]
        if page_entry.get('tab_id'):
            keys.append(page_entry.get('tab_id'))
        if page_entry.get('nested_tab_id'):
            keys.append(page_entry.get('nested_tab_id'))
        if page_entry.get('modal_id'):
            keys.append(page_entry.get('modal_id'))
        return "::".join(filter(None, keys))

    def _validate_manifest(self):
        """
        Validate manifest and generate errors/warnings.
        - missing screenshots
        - missing page descriptions
        - duplicate IDs
        - orphan screenshots
        - orphan descriptions
        - broken image paths
        - invalid manifest schema
        """
        if self.manifest.get("schema_version") != 1:
            self.validation_errors.append("Invalid manifest schema: missing or invalid schema_version")
            
        pages = self.manifest.get("pages", [])
        if not pages:
            self.validation_errors.append("Manifest contains no pages")
            return

        seen_keys = set()
        seen_paths = set()
        
        for p in pages:
            stable_key = self._get_stable_key(p)
            
            # Duplicate IDs
            if stable_key in seen_keys:
                self.validation_warnings.append(f"Duplicate stable key in manifest: {stable_key}")
            seen_keys.add(stable_key)
            
            # Screenshot checks
            s_path = os.path.join(DOCS_DIR, p.get("screenshot_path", ""))
            if not p.get("screenshot_path"):
                self.validation_errors.append(f"Missing screenshot path for entry: {stable_key}")
            elif not os.path.exists(s_path):
                self.validation_errors.append(f"Broken image path/Missing screenshot: {s_path} for key {stable_key}")
            else:
                seen_paths.add(p.get("screenshot_path"))

            # Descriptions check
            desc = self.content.get(stable_key)
            if not desc:
                # Fallback to route
                desc = self.content.get(p.get('route'))
            if not desc:
                self.validation_warnings.append(f"Missing page description for: {stable_key}")

        # Orphan screenshots in screenshots/ folder
        screenshots_dir = os.path.join(DOCS_DIR, "screenshots")
        if os.path.exists(screenshots_dir):
            for file in os.listdir(screenshots_dir):
                if file.endswith(".png"):
                    rel_path = f"screenshots/{file}"
                    if rel_path not in seen_paths:
                        self.validation_warnings.append(f"Orphan screenshot found: {rel_path}")
                        
        # Orphan descriptions
        for content_key in self.content.keys():
            # A description is an orphan if no page in manifest matches its exact key or route
            matched = False
            for p in pages:
                k = self._get_stable_key(p)
                if k == content_key or p.get('route') == content_key:
                    matched = True
                    break
            if not matched:
                self.validation_warnings.append(f"Orphan description entry found: {content_key}")

    def _t(self, key, fallback=None):
        lang_code = self.manifest.get('language', 'en').lower()
        if not hasattr(self, '_i18n_dict'):
            dict_path = os.path.join(BASE_DIR, 'static', 'i18n', f"{lang_code}.json")
            try:
                with open(dict_path, "r", encoding="utf-8") as f:
                    self._i18n_dict = json.load(f)
            except Exception:
                self._i18n_dict = {}
                
        return self._i18n_dict.get(key, fallback if fallback is not None else key)

    def _build_internal_model(self):
        """
        Builds the unified internal document model from manifest and page_descriptions.json
        The order is exactly as dictated by manifest.pages.
        """
        for s in self.manifest.get("pages", []):
            stable_key = self._get_stable_key(s)
            
            desc = self.content.get(stable_key)
            if not desc:
                desc = self.content.get(s.get('route'), {})
            
            raw_title = s.get('modal_id') or s.get('nested_tab_id') or s.get('tab_id') or s.get('page_id') or s.get('route')
            title = str(raw_title).replace('-', ' ').replace('_', ' ').title()
            
            # Use the route or specific id to attempt matching an existing title translation from the frontend keys if available.
            # Usually frontend keys are something like "nav_dashboard" or "tab_dashboard_sett". We will just try a few heuristics 
            # if we wanted to translate the title. However, the requirement is mainly about the headers.
            
            self.internal_model.append({
                "stable_key": stable_key,
                "title": title, # we leave title as is since it's hard to guess the UI key, unless it's in the dict
                "route": s.get('route'),
                "tab": s.get('tab_id'),
                "nested_tab": s.get('nested_tab_id'),
                "modal": s.get('modal_id'),
                "purpose": desc.get("purpose", "[Content pending]"),
                "steps": desc.get("steps", []),
                "screenshot_path": os.path.join(DOCS_DIR, s.get('screenshot_path', '')),
                "relative_screenshot": s.get('screenshot_path', ''),
                "has_content": bool(desc)
            })

    def _filter_model(self, guide_type):
        admin_routes = ["settings", "user-management", "/user-management/", "reports", "advanced-reports"]
        if guide_type == "user":
            return [m for m in self.internal_model if m['route'] not in admin_routes]
        elif guide_type == "admin":
            return [m for m in self.internal_model if m['route'] in admin_routes]
        elif guide_type == "technical":
            return self.internal_model
        return self.internal_model

    def _create_folders(self, guide_type):
        guide_dir = os.path.join(self.output_base_dir, f"{guide_type}_guide")
        for fmt in ["markdown", "html", "pdf", "docx"]:
            os.makedirs(os.path.join(guide_dir, fmt), exist_ok=True)
        return guide_dir

    def _get_filename(self, guide_type, ext):
        lang = self.manifest.get('language', 'EN').upper()
        theme = self.manifest.get('theme', 'Dark').capitalize()
        device = self.manifest.get('device', 'Desktop').capitalize()
        device = "".join(x for x in device if x.isalnum())
        return f"{guide_type.capitalize()}Guide_{lang}_{theme}_{device}_v1.{ext}"

    def generate_all(self):
        os.makedirs(self.output_base_dir, exist_ok=True)
        self._write_validation_report()
        
        for g_type in ["user", "admin", "technical"]:
            self.generate(g_type)
            
        self._update_latest_symlink()

    def _write_validation_report(self):
        report_path = os.path.join(self.output_base_dir, "validation_report.md")
        lines = ["# Documentation Validation Report\n"]
        lines.append(f"**Generated:** {datetime.now().isoformat()}\n")
        
        lines.append(f"## Errors ({len(self.validation_errors)})\n")
        if not self.validation_errors:
            lines.append("No errors.\n")
        for err in self.validation_errors:
            lines.append(f"- [ERROR] {err}")
            
        lines.append(f"\n## Warnings ({len(self.validation_warnings)})\n")
        if not self.validation_warnings:
            lines.append("No warnings.\n")
        for wrn in self.validation_warnings:
            lines.append(f"- [WARN] {wrn}")
            
        with open(report_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _update_latest_symlink(self):
        try:
            if os.path.exists(self.latest_symlink) or os.path.islink(self.latest_symlink):
                if os.path.isdir(self.latest_symlink) and not os.path.islink(self.latest_symlink):
                    shutil.rmtree(self.latest_symlink)
                else:
                    os.unlink(self.latest_symlink)
            
            # Using copytree on Windows to avoid symlink privilege issues, 
            # or directory junction if available. copytree is safer.
            shutil.copytree(self.output_base_dir, self.latest_symlink, dirs_exist_ok=True)
        except Exception as e:
            logger.error(f"Failed to update latest folder: {e}")

    def generate(self, guide_type):
        guide_dir = self._create_folders(guide_type)
        filtered_model = self._filter_model(guide_type)
        if not filtered_model:
            return
            
        md_path = os.path.join(guide_dir, "markdown", self._get_filename(guide_type, "md"))
        html_path = os.path.join(guide_dir, "html", self._get_filename(guide_type, "html"))
        pdf_path = os.path.join(guide_dir, "pdf", self._get_filename(guide_type, "pdf"))
        docx_path = os.path.join(guide_dir, "docx", self._get_filename(guide_type, "docx"))

        self._generate_markdown(filtered_model, md_path, guide_type)
        self._generate_html(md_path, html_path, guide_type)
        self._generate_pdf(html_path, pdf_path)
        self._generate_docx(filtered_model, docx_path, guide_type)

    def _generate_markdown(self, model, out_path, guide_type):
        lang = self.manifest.get('language', 'EN').upper()
        theme = self.manifest.get('theme', 'Dark').capitalize()
        device = self.manifest.get('device', 'Desktop').capitalize()
        app_name = self.manifest.get('application', 'WealthFlow')
        version = self.manifest.get('version', '1.0')
        gen_date = self.manifest.get('generated_at', datetime.now().isoformat())
        
        lines = []
        lines.append("<!-- AUTO-GENERATED START -->\n")
        
        guide_name_key = f"doc_{guide_type}_guide"
        lines.append(f"# {app_name} - {self._t(guide_name_key, guide_type.capitalize() + ' Guide')}\n")
        lines.append(f"**{self._t('doc_version')}:** {version} | **{self._t('doc_language')}:** {lang} | **{self._t('doc_theme')}:** {theme} | **{self._t('doc_device')}:** {device} | **{self._t('doc_date')}:** {gen_date}\n")
        
        # TOC
        lines.append(f"## {self._t('doc_toc')}\n")
        for i, item in enumerate(model, 1):
            anchor = item['title'].lower().replace(' ', '-')
            lines.append(f"{i}. [{item['title']}](#{anchor})")
        lines.append("\n---\n")

        # Content
        for i, item in enumerate(model, 1):
            lines.append(f"## {i}. {item['title']}")
            
            nav_path = " > ".join(filter(None, [item['route'], item['tab'], item['nested_tab'], item['modal']]))
            lines.append(f"**{self._t('doc_navigation')}:** `{nav_path}`\n")
            lines.append(f"**{self._t('doc_purpose')}:** {item['purpose']}\n")
            
            if guide_type != "technical" and item['steps']:
                lines.append(f"**{self._t('doc_steps')}:**")
                for step in item['steps']:
                    lines.append(f"- {step}")
                lines.append("\n")
            elif guide_type == "technical":
                lines.append(f"**{self._t('doc_tech_notes')}:**\n")
                lines.append(f"- {self._t('doc_route')}: `{item['route']}`\n")
                if item['tab']: lines.append(f"- {self._t('doc_tab_id')}: `{item['tab']}`\n")
                if item['modal']: lines.append(f"- {self._t('doc_modal_id')}: `{item['modal']}`\n")
                lines.append("\n")
            
            if os.path.exists(item['screenshot_path']):
                lines.append(f"<figure>")
                lines.append(f"<img src=\"file:///{item['screenshot_path'].replace(chr(92), '/')}\" style=\"max-width: 100%; height: auto; display: block; margin: 0 auto;\" alt=\"{item['title']}\">")
                lines.append(f"<figcaption style=\"text-align:center; font-style:italic;\">{self._t('doc_figure')} {i}: {item['title']}</figcaption>")
                lines.append(f"</figure>\n")
            
            lines.append("---\n")
            
        lines.append("<!-- AUTO-GENERATED END -->\n")
        
        # If writing over an existing markdown file in another location (e.g. repo docs), 
        # but here we generate into a fresh timestamped directory.
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _generate_html(self, md_path, html_path, guide_type):
        app_name = self.manifest.get('application', 'WealthFlow')
        with open(md_path, "r", encoding="utf-8") as f:
            md_text = f.read()
            
        # Strip AUTO-GENERATED markers from HTML
        md_text = md_text.replace("<!-- AUTO-GENERATED START -->\n", "")
        md_text = md_text.replace("<!-- AUTO-GENERATED END -->\n", "")
            
        html_content = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'toc'])
        
        template_path = os.path.join(TEMPLATES_DIR, "html_template.html")
        if os.path.exists(template_path):
            with open(template_path, "r", encoding="utf-8") as f:
                template = f.read()
        else:
            template = "<!DOCTYPE html><html><head><meta charset='utf-8'><title>{{ title }}</title></head><body><div class='container'>{{ content }}</div></body></html>"
            
        guide_name_key = f"doc_{guide_type}_guide"
        final_html = template.replace("{{ title }}", f"{app_name} - {self._t(guide_name_key, guide_type.capitalize() + ' Guide')}").replace("{{ content }}", html_content)
        
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(final_html)

    def _generate_pdf(self, html_path, pdf_path):
        script = os.path.join(DOC_ENGINE_DIR, "html_to_pdf.js")
        if not os.path.exists(script):
            logger.error(f"Cannot generate PDF, missing {script}")
            return
            
        try:
            cmd = ["node", script, html_path, pdf_path]
            subprocess.run(cmd, check=True)
        except subprocess.CalledProcessError as e:
            logger.error(f"PDF generation failed: {e}")
        except Exception as e:
            logger.error(f"Error executing html_to_pdf.js: {e}")

    def _add_page_number(self, run):
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def _generate_docx(self, model, out_path, guide_type):
        doc = Document()
        
        # Headers & Footers
        section = doc.sections[0]
        header = section.header
        header_p = header.paragraphs[0]
        
        guide_name_key = f"doc_{guide_type}_guide"
        header_p.text = f"{self.manifest.get('application', 'WealthFlow')} - {self._t(guide_name_key, guide_type.capitalize() + ' Guide')}"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(self._t('doc_figure', 'Page').replace('الشكل', 'صفحة') + " ") # fallback hack
        self._add_page_number(run)

        # Cover Page
        doc.add_heading(f"{self.manifest.get('application', 'WealthFlow')} - {self._t(guide_name_key, guide_type.capitalize() + ' Guide')}", 0)
        p = doc.add_paragraph()
        p.add_run(f"{self._t('doc_version')}: {self.manifest.get('version', '1.0')}\n")
        p.add_run(f"{self._t('doc_language')}: {self.manifest.get('language', 'EN').upper()}\n")
        p.add_run(f"{self._t('doc_theme')}: {self.manifest.get('theme', 'Dark').capitalize()}\n")
        p.add_run(f"{self._t('doc_device')}: {self.manifest.get('device', 'Desktop').capitalize()}\n")
        p.add_run(f"{self._t('doc_date')}: {self.manifest.get('generated_at', datetime.now().isoformat())}")
        doc.add_page_break()
        
        # TOC
        doc.add_heading(self._t('doc_toc'), level=1)
        for i, item in enumerate(model, 1):
            doc.add_paragraph(f"{i}. {item['title']}", style='List Number')
            
        doc.add_page_break()
        
        # Content
        for i, item in enumerate(model, 1):
            doc.add_heading(f"{i}. {item['title']}", level=2)
            
            nav_path = " > ".join(filter(None, [item['route'], item['tab'], item['nested_tab'], item['modal']]))
            p = doc.add_paragraph()
            p.add_run(f"{self._t('doc_navigation')}: ").bold = True
            p.add_run(nav_path)
            
            p2 = doc.add_paragraph()
            p2.add_run(f"{self._t('doc_purpose')}: ").bold = True
            p2.add_run(item['purpose'])
            
            if guide_type != "technical" and item['steps']:
                doc.add_paragraph(self._t("doc_steps") + ":", style='Heading 3')
                for step in item['steps']:
                    doc.add_paragraph(step, style='List Bullet')
            elif guide_type == "technical":
                doc.add_paragraph(self._t("doc_tech_notes") + ":", style='Heading 3')
                doc.add_paragraph(f"{self._t('doc_route')}: {item['route']}", style='List Bullet')
                if item['tab']: doc.add_paragraph(f"{self._t('doc_tab_id')}: {item['tab']}", style='List Bullet')
                if item['modal']: doc.add_paragraph(f"{self._t('doc_modal_id')}: {item['modal']}", style='List Bullet')
            
            if os.path.exists(item['screenshot_path']):
                try:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_img = p_img.add_run()
                    r_img.add_picture(item['screenshot_path'], width=Inches(6.0))
                    
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_cap = p_cap.add_run(f"{self._t('doc_figure')} {i}: {item['title']}")
                    r_cap.italic = True
                except Exception as e:
                    logger.warning(f"Could not add image {item['screenshot_path']} to docx: {e}")
            
            if i < len(model):
                doc.add_page_break()
            
        doc.save(out_path)
